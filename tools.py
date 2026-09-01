from docx import Document
from docx.shared import Pt
import os


def create_word_document(user_request, plan, content):
    """
    Creates a professional Word document.
    Uses /tmp for serverless environments such as Vercel.
    """

    document = Document()

    # -------------------------
    # Title
    # -------------------------
    title = document.add_heading(
        "AI Generated Business Document",
        level=1
    )
    title.style.font.size = Pt(18)

    # -------------------------
    # User Request
    # -------------------------
    document.add_heading("User Request", level=2)
    document.add_paragraph(user_request)

    # -------------------------
    # Execution Plan
    # -------------------------
    document.add_heading("Agent Execution Plan", level=2)

    for step in plan:
        document.add_paragraph(step, style="List Bullet")

    # -------------------------
    # AI Generated Document
    # -------------------------
    document.add_heading("Generated Document", level=2)

    lines = content.split("\n")

    headings = [
        "Title",
        "Executive Summary",
        "Introduction",
        "Problem Statement",
        "Solution",
        "Features",
        "Benefits",
        "Implementation Plan",
        "Financials",
        "Pricing",
        "Conclusion",
        "Call to Action"
    ]

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # Convert known sections into Word headings
        if any(line.lower().startswith(h.lower()) for h in headings):
            document.add_heading(line, level=3)

        # Preserve bullet points
        elif line.startswith("-") or line.startswith("*"):
            document.add_paragraph(
                line[1:].strip(),
                style="List Bullet"
            )

        else:
            document.add_paragraph(line)

    # -------------------------
    # Save document
    # -------------------------
    if os.getenv("VERCEL"):
        output_file = "/tmp/output.docx"
    else:
        output_file = "output.docx"

    document.save(output_file)

    return output_file